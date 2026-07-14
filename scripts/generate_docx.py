import os
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = docx.Document()

    # Set page margins (1 inch on all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Theme colors (NeoConta Dark Mode & Accent palette)
    COLOR_PURPLE = RGBColor(26, 11, 46)     # Primary dark brand color (#1A0B2E)
    COLOR_ORANGE = RGBColor(255, 94, 0)     # Highlight color (#FF5E00)
    COLOR_GRAY = RGBColor(74, 85, 104)      # Body text color (#4A5568)
    COLOR_LIGHT_GRAY = RGBColor(113, 128, 150) # Secondary body text
    COLOR_BLACK = RGBColor(22, 11, 36)

    # Styles helper functions
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = COLOR_PURPLE
        p.paragraph_format.space_after = Pt(4)
        
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run("Compilación de contenidos en orden de aparición")
        run_sub.font.name = 'Arial'
        run_sub.font.size = Pt(11)
        run_sub.font.italic = True
        run_sub.font.color.rgb = COLOR_LIGHT_GRAY
        p_sub.paragraph_format.space_after = Pt(24)

    def add_section_header(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = COLOR_PURPLE
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True

    def add_subheading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = COLOR_ORANGE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True

    def add_body(text, bold=False, italic=False, color=COLOR_GRAY, before=0, after=6):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        return p

    def add_bullet(bold_prefix, text_content, color=COLOR_GRAY):
        p = doc.add_paragraph(style='List Bullet')
        if bold_prefix:
            run_pre = p.add_run(bold_prefix)
            run_pre.font.name = 'Arial'
            run_pre.font.size = Pt(10.5)
            run_pre.font.bold = True
            run_pre.font.color.rgb = COLOR_BLACK
        run_text = p.add_run(text_content)
        run_text.font.name = 'Arial'
        run_text.font.size = Pt(10.5)
        run_text.font.color.rgb = color
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)

    # -------------------------------------------------------------
    # Document Construction
    # -------------------------------------------------------------
    add_title("NeoConta - Textos Oficiales de la Web")

    # --- NAVBAR ---
    add_section_header("1. Barra de Navegación (Navbar)")
    add_body("La barra de navegación es visible en la parte superior de la página principal (se muestra al hacer scroll).")
    add_bullet("Logo: ", "NeoConta (Carga el archivo visual '/assets/navbar_logo.png')")
    add_bullet("Enlace 1: ", "Inicio")
    add_bullet("Enlace 2: ", "Funcionalidades")
    add_bullet("Enlace 3: ", "Contacto")
    add_bullet("Botón secundario: ", "Registrarse")
    add_bullet("Botón principal (CTA): ", "Iniciar Sesión (con ícono de flecha hacia la derecha)")

    # --- HERO PRINCIPAL ---
    add_section_header("2. Sección Hero de Bienvenida (Pantalla Completa)")
    add_body("Esta es la sección de impacto visual que el usuario ve al cargar el sitio por primera vez.")
    add_bullet("Imagen de Fondo: ", "Fondo tecnológico/futurista (Carga el archivo visual '/assets/hero-bg.jpg')")
    add_bullet("Capa de contraste: ", "Overlay oscuro semitransparente (bg-black/20)")
    add_bullet("Logotipo Centrado: ", "NeoConta en color blanco (Carga el archivo visual '/assets/logo-white.png')")
    add_bullet("Indicador de Desplazamiento: ", "Ícono animado de mouse al final de la pantalla que invita a deslizar hacia abajo.")

    # --- NUESTRA MISION ---
    add_section_header("3. Sección: Nuestra Misión (Original Hero)")
    add_body("Sección de introducción y propuesta de valor de la compañía, ubicada justo debajo del Hero de bienvenida.")
    add_bullet("Etiqueta destacada: ", "NUESTRA MISIÓN")
    add_bullet("Título principal: ", "Transformación Inteligente")
    add_bullet("Texto descriptivo: ", 
               "Integramos la experiencia en gestión de negocios con consultorías especializadas en Inteligencia Artificial. "
               "Transformamos la realidad de emprendedores y PyMES aplicando IA en áreas clave: desde la automatización de "
               "procesos operativos y financieros para lograr eficiencia escalable, hasta el desarrollo de estrategias de "
               "contenido y comunicación impulsadas por inteligencia artificial.")
    add_bullet("Elemento Visual (Derecha): ", "Imagen conceptual de red neuronal (Carga el archivo visual '/assets/ai-network.jpg') con brillos en degradé naranja y violeta.")

    # --- FUNCIONALIDADES ---
    add_section_header("4. Sección: Funcionalidades")
    add_body("Eje central de la página donde se detallan las 4 áreas o soluciones clave provistas por NeoConta.")
    add_subheading("Encabezado de la Sección")
    add_bullet("Título: ", "El poder de la IA potenciando la Gestión de Negocios")
    add_bullet("Descripción: ", 
               "Implementamos soluciones de IA que potencian cada negocio eliminando tareas manuales repetitivas, "
               "dando espacio a aquellas con verdadero valor personal. Estructuramos la información y la convertimos "
               "en una auténtica ingeniería de datos, optimizando así la gestión integral de proyectos y organizaciones.")
    
    add_subheading("Tarjetas de Funcionalidad (Tarjetas Informativas)")
    
    add_bullet("Tarjeta 1 - Estrategia de IA para Negocios: ", 
               "Diagnosticamos tu empresa, tus procesos y tus objetivos para diseñar un plan concreto de incorporación "
               "de inteligencia artificial. Nada de humo futurista: identificamos dónde la IA puede ahorrar tiempo, "
               "reducir errores, mejorar decisiones y generar nuevas oportunidades de crecimiento. Resultado: una hoja "
               "de ruta clara para aplicar IA de forma práctica, rentable y escalable.")
    add_bullet("Tarjeta 2 - Automatización Inteligente: ", 
               "Diseñamos flujos automatizados con IA para reducir tareas repetitivas, conectar herramientas y optimizar "
               "la gestión diaria de tu organización. Desde carga de datos y reportes hasta respuestas automáticas, "
               "tableros, alertas y asistentes internos. Resultado: menos trabajo manual, más velocidad operativa y "
               "equipos enfocados en lo importante.")
    add_bullet("Tarjeta 3 - Commander©: ", 
               "Plataforma de Business Intelligence personalizada. Sin importar el tamaño de tu empresa, con Commander© "
               "tendrás toda la información clave al alcance de tu mano para dominar cada decisión de tu proyecto.")
    add_bullet("Tarjeta 4 - Contenido y Cultura IA: ", 
               "Acompañamos a empresas, profesionales y equipos en el aprendizaje real de la inteligencia artificial. "
               "Diseñamos capacitaciones, guías de uso, contenidos estratégicos y metodologías para que la IA no sea una moda, "
               "sino una herramienta incorporada con criterio. Resultado: equipos más preparados, más autónomos y con menos "
               "miedo a usar tecnología.")

    # --- METRICAS ---
    add_section_header("5. Sección: Métricas y Logros (Estadísticas)")
    add_body("Banda de datos cuantitativos que valida la trayectoria de la firma.")
    add_bullet("Métrica 1: ", "Clientes Acompañados: 10+")
    add_bullet("Métrica 2: ", "Procesos Repensados: 30+")
    add_bullet("Métrica 3: ", "Soluciones Brindadas: 15+")
    add_bullet("Métrica 4: ", "Personalización: 100%")

    # --- FORMULARIO DE CONTACTO ---
    add_section_header("6. Sección: Formulario de Contacto")
    add_body("Sección interactiva para la captación de prospectos y consultas generales.")
    add_bullet("Título principal: ", "¿Hablamos de tu proyecto?")
    add_bullet("Descripción: ", "Dejanos tu consulta y un especialista en Inteligencia Artificial y gestión se pondrá en contacto con vos a la brevedad.")
    
    add_subheading("Campos y Etiquetas del Formulario (Labels & Placeholders)")
    add_bullet("Nombre Completo: ", "Etiqueta: 'Nombre completo' | Sugerencia (Placeholder): 'Juan Pérez' (Obligatorio)")
    add_bullet("Correo Electrónico: ", "Etiqueta: 'Correo electrónico' | Sugerencia (Placeholder): 'juan@ejemplo.com' (Obligatorio)")
    add_bullet("Celular/Teléfono: ", "Etiqueta: 'Celular / Teléfono' | Sugerencia (Placeholder): '+54 9 11 1234 5678' (Opcional)")
    add_bullet("Mensaje: ", "Etiqueta: 'Mensaje' | Sugerencia (Placeholder): 'Hola, me gustaría saber más sobre cómo implementar IA en mi empresa...' (Obligatorio)")
    add_bullet("Botón de Acción: ", "Enviar mensaje (En reposo) | Enviando... (Cargando)")
    
    add_subheading("Mensajes de Estado (Success & Error States)")
    add_bullet("Mensaje de Éxito: ", "¡Mensaje Enviado! Gracias por contactarnos. Tu mensaje ha sido recibido con éxito y nos comunicaremos con vos a la brevedad. [Botón: Enviar otro mensaje]")
    add_bullet("Mensaje de Error 1 (Servidor): ", "Ocurrió un error inesperado. (O el mensaje enviado por la API)")
    add_bullet("Mensaje de Error 2 (Red): ", "No se pudo conectar con el servidor. Verifique su conexión.")

    # --- PIE DE PAGINA ---
    add_section_header("7. Pie de Página (Footer)")
    add_body("Ubicado de forma fija al pie del sitio web con información legal y enlaces rápidos.")
    add_bullet("Nombre de Marca: ", "NeoConta")
    add_bullet("Línea descriptiva: ", "Redefiniendo tu administración de la mano de la Inteligencia Artificial.")
    add_bullet("Redes Sociales: ", "Twitter/X, LinkedIn, GitHub (Redirección a perfiles oficiales)")
    add_bullet("Columna Compañía: ", "Acerca de (Abre Modal) | Carreras (Abre Modal) | Contacto (Enlace)")
    add_bullet("Columna Legal: ", "Privacidad (Abre Modal) | Términos (Abre Modal)")
    add_bullet("Derechos de Autor: ", "© 2026 NeoConta. Todos los derechos reservados. (El año se renderiza de forma dinámica)")
    add_bullet("Créditos: ", "Hecho con [Corazón] en Argentina")

    # --- MODALES EMERGENTES ---
    add_section_header("8. Contenidos de los Modales del Footer (Emergentes)")
    add_body("Estos textos se muestran en ventanas emergentes (modales) al hacer clic en los enlaces del Footer.")
    
    # Modal Acerca De
    add_subheading("A) Modal: Acerca de (Sobre NeoConta)")
    add_bullet("Título del Modal: ", "Sobre NeoConta")
    add_body("En NeoConta redefinimos la administración empresarial a través de la consultoría estratégica y la "
             "integración de Inteligencia Artificial. Nos especializamos en repensar y automatizar los procesos "
             "contables y operativos de tu negocio, transformando la gestión administrativa tradicional en un "
             "ecosistema ágil, eficiente y adaptado a la era digital.", color=COLOR_GRAY)
    add_bullet("Acción: ", "Cerrar")

    # Modal Carreras
    add_subheading("B) Modal: Carreras (Sumate al Equipo)")
    add_bullet("Título del Modal: ", "Sumate al Equipo")
    add_bullet("Subtítulo: ", "Buscamos profesionales proactivos y apasionados por redefinir procesos contables y administrativos de la mano de la Inteligencia Artificial.")
    add_bullet("Campos de Formulario: ", 
               "- Nombre Completo * (Sugerencia: Tu nombre y apellido)\n"
               "- Correo Electrónico * (Sugerencia: ejemplo@correo.com)\n"
               "- Área de Interés * (Opciones seleccionables: Consultoría de Procesos, Desarrollo de Software / IA, Administración y Finanzas, Soporte y Operaciones)\n"
               "- Cargar CV (PDF, DOC o DOCX) * (Texto: Arrastrá tu archivo aquí o hacé clic para buscar | Límite de 5MB)")
    add_bullet("Botones de Acción: ", "Cancelar | Enviar")
    add_bullet("Estado Exitoso: ", "¡Postulación Recibida! Gracias por tu interés en sumarte a NeoConta. Evaluaremos tu perfil y nos pondremos en contacto contigo a la brevedad. [Botón: Entendido]")
    add_bullet("Mensajes de Validación y Error: ", 
               "- El nombre es obligatorio.\n"
               "- Por favor ingrese un correo válido.\n"
               "- Debe adjuntar su CV.\n"
               "- El archivo debe pesar menos de 5MB.\n"
               "- Solo se permiten archivos PDF, DOC o DOCX.\n"
               "- Error de conexión. Intente de nuevo más tarde.")

    # Modal Privacidad
    add_subheading("C) Modal: Declaración de Privacidad")
    add_bullet("Título del Modal: ", "Declaración de Privacidad")
    add_body("En NeoConta asumimos un compromiso absoluto con la transparencia, la seguridad de la información y el cumplimiento del marco regulatorio de la República Argentina.")
    
    add_body("1. Protección de Datos Personales (Ley N° 25.326)", bold=True, color=COLOR_BLACK)
    add_body("De conformidad con lo dispuesto por la Ley N° 25.326 de Protección de Datos Personales, sus normas reglamentarias y complementarias, le informamos que:")
    
    add_bullet("Consentimiento Informado: ", "Al ingresar sus datos en nuestros formularios o cargar su currículum vitae en la sección de Carreras, usted otorga su consentimiento previo, libre e informado (Art. 5 y 6) para el tratamiento de su información.")
    add_bullet("Finalidad del Tratamiento: ", "Los datos recopilados serán utilizados exclusivamente para los fines para los cuales fueron suministrados (evaluación de perfiles laborales, contacto institucional o provisión de servicios de consultoría administrativa y tecnológica).")
    add_bullet("Deber de Seguridad y Confidencialidad: ", "Implementamos estrictas medidas técnicas y organizativas para garantizar la seguridad de la información (Art. 9 de la Ley 25.326), evitando su pérdida, alteración o el acceso no autorizado por parte de terceros. Asimismo, mantenemos secreto profesional sobre sus datos (Art. 10).")
    add_bullet("Derechos ARCO: ", "Como titular de los datos, usted tiene el derecho legal de acceder de forma gratuita a sus datos personales (Art. 14), así como de solicitar su rectificación, actualización o supresión (Art. 16) mediante comunicación directa a nuestros canales oficiales de contacto.")
    
    add_body("2. Propiedad Intelectual y Limitación de Responsabilidad (Ley N° 22.362 y CCyCN)", bold=True, color=COLOR_BLACK)
    add_bullet("Registro Marcario: ", "NeoConta es titular exclusivo de los derechos de propiedad intelectual, marcas, logotipos y nombres comerciales asociados bajo los términos de la Ley de Marcas y Designaciones N° 22.362.")
    add_bullet("Uso no Autorizado: ", "Queda estrictamente prohibido el uso indebido de nuestra marca \"NeoConta\", nombre comercial, logotipos o cualquier otro signo identificatorio por parte de terceros no autorizados a su uso, incluyendo cualquier reproducción, imitación o explotación no autorizada. Cualquier infracción será pasible de las acciones y sanciones civiles y penales previstas en el Art. 31 de la Ley 22.362.")
    add_bullet("Deslinde de Responsabilidad: ", "Garantizamos el resguardo técnico y la seguridad de las plataformas digitales bajo nuestro control directo. Sin embargo, NeoConta declina expresamente toda responsabilidad por los daños y perjuicios de cualquier naturaleza que puedan derivarse del uso ilegal de nuestro nombre y marca por parte de terceros no autorizados (incluyendo ataques de suplantación de identidad o phishing, páginas web apócrifas, perfiles falsos en redes sociales o correos fraudulentos ajenos a nuestra organización). Instamos a los usuarios a constatar la legitimidad de toda comunicación mediante nuestros canales oficiales.")
    
    add_body("Última actualización: Mayo de 2026.", italic=True, color=COLOR_LIGHT_GRAY)

    # Modal Terminos
    add_subheading("D) Modal: Términos y Condiciones")
    add_bullet("Título del Modal: ", "Términos y Condiciones")
    add_body("Bienvenido a NeoConta. Al acceder, navegar o utilizar nuestra plataforma digital y servicios de consultoría, usted acepta cumplir y estar sujeto a los siguientes Términos y Condiciones.")
    
    add_body("1. Aceptación y Encuadre Legal (Arts. 984 a 989 del CCyCN)", bold=True, color=COLOR_BLACK)
    add_body("Estos Términos y Condiciones constituyen un contrato de adhesión celebrado por medios electrónicos. Al interactuar con el sitio web o utilizar las herramientas de NeoConta, usted presta su libre conformidad con la totalidad de estas cláusulas, de acuerdo con los artículos 984, 985, 986 y concordantes del Código Civil y Comercial de la Nación Argentina.")
    
    add_body("2. Naturaleza y Alcance de los Servicios (Ley N° 24.240)", bold=True, color=COLOR_BLACK)
    add_bullet("Objeto: ", "NeoConta brinda servicios de consultoría estratégica y desarrollo de soluciones basadas en Inteligencia Artificial para la optimización de procesos contables y administrativos.")
    add_bullet("Ausencia de Asesoramiento Matricular Directo: ", "La plataforma y sus reportes actúan como herramientas de asistencia y automatización. No constituyen, en sí mismas, un reemplazo del dictamen profesional de contadores públicos matriculados según las normativas de los consejos profesionales locales.")
    add_bullet("Deber de Información (Art. 4 de la Ley N° 24.240): ", "Nos comprometemos a proveer de forma clara, detallada y gratuita toda la información sobre el funcionamiento de las herramientas y servicios brindados.")
    
    add_body("3. Contratación por Medios Electrónicos (Arts. 1105 a 1116 del CCyCN)", bold=True, color=COLOR_BLACK)
    add_bullet("Perfeccionamiento: ", "Los acuerdos y solicitudes de contacto/carrera se perfeccionan a distancia mediante medios electrónicos.")
    add_bullet("Notificaciones: ", "Todas las comunicaciones enviadas al correo electrónico consignado por el usuario o a través de la plataforma serán consideradas válidas y vinculantes.")
    
    add_body("4. Responsabilidades y Limitaciones de Servicio", bold=True, color=COLOR_BLACK)
    add_bullet("Exactitud de los Datos: ", "El usuario es exclusivamente responsable por la veracidad, licitud y exactitud de la información, facturas, CVs o datos bancarios que cargue o integre en el sistema.")
    add_bullet("Disponibilidad del Servicio: ", "Garantizamos la máxima diligencia técnica para el mantenimiento del sitio. Sin embargo, NeoConta no será responsable por interrupciones, demoras o errores causados por fallas en servicios externos ajenos a nuestro control directo (por ejemplo, interrupciones en los sistemas de AFIP/ARCA, plataformas bancarias integradas o fallas de red del proveedor de internet del usuario).")
    add_bullet("Abuso y Cláusulas Nulas: ", "En cumplimiento del artículo 1118 y concordantes del CCyCN y la Ley N° 24.240, cualquier cláusula de este documento que sea declarada nula o abusiva por la autoridad competente no afectará la validez del resto del acuerdo.")
    
    add_body("5. Ley Aplicable y Jurisdicción", bold=True, color=COLOR_BLACK)
    add_body("Este acuerdo se rige por las leyes vigentes de la República Argentina. Para cualquier disputa que no pueda resolverse amigablemente, las partes se someten a la jurisdicción de los tribunales competentes en materia comercial de la República Argentina, salvaguardando los derechos de prórroga de jurisdicción y domicilio del consumidor establecidos en la Ley N° 24.240.")
    
    add_body("Última actualización: Mayo de 2026.", italic=True, color=COLOR_LIGHT_GRAY)

    # --- APENDICE ---
    add_section_header("9. Apéndice: Textos Secundarios / Legados (Unused Components)")
    add_body("Estos textos forman parte de la base de código en componentes alternativos no vinculados a la página principal activa actualmente.")
    
    add_subheading("Componente: Hero alternativo (Hero.js)")
    add_bullet("Título: ", "NeoConta")
    add_bullet("Subtítulo: ", "Soluciones contables digitales para el futuro de tu negocio.")
    add_bullet("Botón: ", "Descubrir Servicios")
    
    add_subheading("Componente: Servicios alternativos (Services.js)")
    add_bullet("Título Sección: ", "Nuestros Servicios")
    add_bullet("Descripción Sección: ", "Hacemos que la contabilidad sea simple, moderna y eficiente.")
    add_bullet("Servicio 1: ", "Facturador Masivo AFIP | Automatiza la emisión de facturas electrónicas. Conexión directa y segura con AFIP.")
    add_bullet("Servicio 2: ", "Consultoría Financiera | Análisis detallado de tus finanzas para tomar las mejores decisiones de negocio.")
    add_bullet("Servicio 3: ", "Gestión de Nómina | Liquidación de sueldos y cargas sociales al día, sin complicaciones.")
    add_bullet("Servicio 4: ", "Asesoría Impositiva | Optimización fiscal y cumplimiento de todas tus obligaciones tributarias.")

    # Save document
    output_path = "textos_pagina_principal.docx"
    doc.save(output_path)
    print(f"Documento creado exitosamente en: {os.path.abspath(output_path)}")

if __name__ == "__main__":
    create_document()
